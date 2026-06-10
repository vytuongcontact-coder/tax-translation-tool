import os
import json
import re
import fitz  # PyMuPDF
from docx import Document

# ── 1. TẢI TỪ ĐIỂN VÀ KHỞI TẠO BIỂU THỨC CHÍNH QUY (REGEX) ──
dict_path = os.path.join(os.path.dirname(__file__), "tax_dictionary.json")

pattern = None

if os.path.exists(dict_path):
    with open(dict_path, 'r', encoding='utf-8') as f:
        # Chuẩn hóa dấu nháy và lọc bỏ các từ chỉ có 1 âm tiết để tránh nhiễu
        terms = set(t.replace("’", "'").strip() for t in json.load(f))
        filtered_terms = [t for t in terms if len(t.split()) >= 2]
        
        if filtered_terms:
            # Sắp xếp thuật ngữ từ dài nhất đến ngắn nhất để regex quét tham lam (greedy) chính xác nhất
            sorted_terms = sorted(filtered_terms, key=len, reverse=True)
            # Escape các ký tự đặc biệt trong từ khóa
            escaped_terms = [re.escape(t) for t in sorted_terms]
            # Tạo regex pattern ghép các từ khóa bằng toán tử HOẶC | và bao bọc bởi ranh giới từ \b
            pattern_str = r'\b(' + '|'.join(escaped_terms) + r')\b'
            pattern = re.compile(pattern_str, re.IGNORECASE)

def extract_terms(file_path: str) -> list:
    """Quét thuật ngữ sử dụng biểu thức chính quy (Regex) tối ưu hóa hỗ trợ cả Word (.docx) và PDF (.pdf)"""
    
    # ── 2. TRÍCH XUẤT VÀ CHUẨN HÓA VĂN BẢN ──
    texts = []
    
    if file_path.lower().endswith('.pdf'):
        # Trích xuất văn bản từ tệp PDF bằng PyMuPDF
        with fitz.open(file_path) as doc:
            for page in doc:
                text = page.get_text()
                if text.strip():
                    # Xử lý ngắt dòng giữa các từ trong trang PDF
                    clean_text = text.replace('\n', ' ')
                    texts.append(clean_text)
    else:
        # Trích xuất văn bản từ tệp Word (.docx) bằng python-docx
        doc = Document(file_path)
        # Lấy text từ các đoạn văn
        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text.strip())
                
        # Lấy text từ bảng biểu
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        # Xử lý lỗi ngắt dòng giữa chừng trong bảng (VD: "Arm's \n Length")
                        clean_cell = cell.text.strip().replace('\n', ' ')
                        texts.append(clean_cell)
                        
    # Tạo thành một chuỗi duy nhất để quét
    full_text = " . ".join(texts)
    # Chuẩn hóa dấu nháy trong toàn bộ văn bản để khớp chính xác với từ điển
    full_text = full_text.replace("’", "'").replace("‘", "'")
    
    # ── 3. QUÉT TỐC ĐỘ CAO BẰNG REGEX ──
    candidate_terms = set()
    if pattern:
        for m in pattern.finditer(full_text):
            # Lấy nguyên bản định dạng chữ Hoa/Thường (casing) từ file tài liệu
            candidate_terms.add(m.group(0))
                
    # Sắp xếp theo độ dài giảm dần, và ưu tiên chữ thường nếu độ dài bằng nhau để quét tham lam tốt nhất
    sorted_terms = sorted(list(candidate_terms), key=lambda x: (len(x), x.islower()), reverse=True)
    
    # Loại bỏ trùng lặp không phân biệt hoa thường (Case-insensitive deduplication)
    seen_lower = set()
    unique_terms = []
    for term in sorted_terms:
        term_lower = term.lower()
        if term_lower not in seen_lower:
            seen_lower.add(term_lower)
            unique_terms.append(term)
            
    return unique_terms

def extract_raw_text(file_path: str) -> str:
    """Extract raw text from PDF or DOCX for the frontend preview pane"""
    texts = []
    if file_path.lower().endswith('.pdf'):
        with fitz.open(file_path) as doc:
            for page in doc:
                text = page.get_text()
                if text.strip():
                    texts.append(text)
    else:
        doc = Document(file_path)
        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)
    full_text = "\n\n".join(texts)
    # Normalize curly/smart quotes to straight quotes for exact matching in UI
    return full_text.replace("’", "'").replace("‘", "'").replace("“", '"').replace("”", '"')
